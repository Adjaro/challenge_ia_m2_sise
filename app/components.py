import os
import streamlit as st
from PIL import Image  # Pour gérer l'image du logo
import json
from docx import Document
from io import BytesIO

from utils.ia import (
    analyze_offre_emploi,
    calculate_section_similarity,
    generate_lettre_motivation,
    calculate_similarity,
)
from utils.classFactory import ScrapingFactory
import time
from utils.monitoring_ecologie import EnvironmentMetrics

LOGO_PATH = "logo.png"


def set_custom_style():
    """
    Définit le style CSS personnalisé pour la sidebar et les boutons.
    """
    st.markdown(
        """
        <style>
            /* Style pour la sidebar */
            section[data-testid="stSidebar"] {
                background-color: #3f4f43 !important; /* Couleur de fond de la sidebar */
                border-radius: 20px !important; /* Arrondir les coins de la sidebar */
                padding: 20px !important; /* Espacement interne */
                color: white !important; /* Couleur du texte dans la sidebar */
            }

            /* Style pour les boutons dans la sidebar */
            .stButton > button {
                border-radius: 10px !important;
                padding: 10px 20px !important;
                font-size: 16px !important;
                transition: all 0.3s ease !important;
                background-color: #ffffff !important; /* Couleur de fond du bouton */
                color: #45a049 !important; /* Couleur du texte du bouton */
                border: 2px solid #45a049 !important; /* Bordure verte */
                width: 100% !important; /* Largeur maximale */
                margin: 10px 0 !important; /* Marge entre les boutons */
            }

            .stButton > button:hover {
                background-color: #45a049 !important; /* Couleur de fond au survol */
                color: white !important; /* Couleur du texte au survol */
                transform: scale(1.05) !important; /* Effet de zoom au survol */
            }
 
        </style>
    """,
        unsafe_allow_html=True,
    )


def load_logo():
    """
    Charge et affiche le logo dans la sidebar.
    """
    st.sidebar.image("logo.png", width=250)  # Remplacez par le chemin de votre logo


def show_modifier_cv():
    """
    Affiche la boîte de dialogue pour modifier le CV.
    """
    st.success("Fonctionnalité de modification du CV")


def show_sidebar(cv_info: dict) -> str:
    """
    Affiche la barre latérale de l'application avec les informations du CV.

    Args:
        cv_info (dict): Dictionnaire contenant les informations du CV.

    Returns:
        str: Nom de la conversation sélectionnée.
    """
    set_custom_style()
    load_logo()  # Charger le logo

    with st.sidebar:
        st.markdown("---")

        # Button "Modifier mes informations"
        if st.button("Modifier mes informations"):
            show_modifier_cv()

        # Button "Voir plus d'offres"
        if st.button("Voir plus d'offres"):
            st.session_state["domaine_page"] = 1

        if st.session_state.get("domaine_page") == 1:
            entrer_domain()
        elif st.session_state.get("domaine_page") == 2:
            show_offre_emploi()

        # Button "Impact environnemental"
        if st.button("Impact environnemental"):
            class_impact = st.session_state["monitoring_environnement"]
            gwp, energy_usage = class_impact.get_metrics()
            st.success(
                f"Impact environnemental:\nGWP: {gwp} kgCO2eq \n Energy Usage: {energy_usage} kWh"
            )

        st.markdown("---")


def comparer_cv_offre():
    analyse = st.session_state.get("offre_json", {})
    str_analyse = json.dumps(analyse)

    cv_json = st.session_state.get("cv_json", {})
    str_cv = json.dumps(cv_json)

    comparaison = calculate_section_similarity(
        cv_reformule=str_cv, offre_emploi_reformule=str_analyse
    )
    st.session_state["comparaison"] = comparaison
    return comparaison


def comparer_cv():
    """
    Affiche une interface pour comparer un CV à une offre d'emploi.
    """
    st.container(height=200, border=False)
    # Conteneur pour l'en-tête
    with st.container(border=False):
        st.markdown("### 📄 Comparer votre CV à une offre d'emploi")
        st.markdown(
            "Saisissez l'URL de l'offre d'emploi et posez une question pour obtenir une comparaison détaillée."
        )

    # Champ de saisie pour l'URL de l'offre d'emploi
    # st.text_area(label, value="", height=None, max_chars=None, key=None, help=None, on_change=None, args=None, kwargs=None, *, placeholder=None, disabled=False, label_visibility="visible")

    url = st.text_area(
        label="URL de l'offre d'emploi",
        placeholder="https://www.example.com",
        help="Veuillez saisir l'URL de l'offre d'emploi à comparer.",
        height=70,
    )

    # Bouton pour lancer la comparaison
    if st.button("Comparer", key="compare_button"):
        if url.strip() == "" or url == "https://www.example.com":
            st.warning("Veuillez saisir une URL valide.")
        elif (
            "candidat.francetravail.fr" not in url.lower()
        ):  # Vérifier que l'URL contient "france travail"
            st.error(
                "L'URL doit provenir de France Travail. Veuillez saisir une URL valide."
            )
        else:

            with st.spinner("Comparaison en cours...", show_time=True):
                # time.sleep(5)
                progress_bar = st.progress(0)
                scraper = ScrapingFactory()

                progress_bar.progress(10)
                offre_emploi = scraper.scrap_one(url)
                st.session_state["offre_emploi_brut"] = offre_emploi
                # st.write(offre_emploi)

                progress_bar.progress(30)
                analyse = analyze_offre_emploi(offre_emploi)
                # st.write(analyse)
                analyse = json.loads(analyse)
                st.session_state["offre_json"] = analyse

                # st.write(analyse)

                progress_bar.progress(60)
                comparaison = comparer_cv_offre()

                progress_bar.progress(100)
                st.success("La comparaison a été effectuée avec succès !")

                show_comparaison_cv()

                if offre_emploi is None:
                    st.error(
                        "Impossible de récupérer les informations de l'offre d'emploi. Veuillez réessayer."
                    )
                    return
                else:
                    st.write(comparaison)


@st.dialog("Comparaison du CV", width="large")
def show_comparaison_cv():
    """
    Affiche un dialogue pour comparer le CV à une offre d'emploi avec un design amélioré.
    """
    # Custom CSS for the comparison display
    st.markdown(
        """
        <style>
            .comparison-header {
                font-size: 24px;
                font-weight: bold;
                margin-bottom: 20px;
                color: #1E3D59;
            }
            .score-card {
                padding: 20px;
                border-radius: 10px;
                background-color: #f8f9fa;
                margin-bottom: 15px;
            }
            .score-title {
                font-size: 18px;
                font-weight: bold;
                color: #2C3E50;
            }
            .match-percentage {
                font-size: 32px;
                font-weight: bold;
                margin: 10px 0;
            }
            .section-content {
                margin-top: 10px;
                padding: 15px;
                background-color: white;
                border-radius: 8px;
                border-left: 4px solid #4CAF50;
            }
        </style>
    """,
        unsafe_allow_html=True,
    )

    comparaison = st.session_state.get("comparaison", {})
    cv_data = st.session_state.get("cv_json", {})
    offre_data = st.session_state.get("offre_json", {})

    st.markdown(
        '<p class="comparison-header">🎯 Résultats de la comparaison</p>',
        unsafe_allow_html=True,
    )

    # Create tabs for different views
    tab1, tab2 = st.tabs(["📊 Vue d'ensemble", "📑 Détails par section"])

    with tab1:
        # Overview with metrics
        col1, col2, col3, col4 = st.columns(4)

        metrics = {
            "Formation": (comparaison.get("formation", 0) * 100, "🎓"),
            "Compétences": (comparaison.get("competences", 0) * 100, "💪"),
            "Expériences": (comparaison.get("experiences", 0) * 100, "💼"),
            "Profil": (comparaison.get("profil", 0) * 100, "👤"),
        }

        for (label, (value, emoji)), col in zip(
            metrics.items(), [col1, col2, col3, col4]
        ):
            with col:
                st.metric(
                    label=f"{emoji} {label}", value=f"{value:.1f}%", delta="match"
                )
        col1, col2, col3 = st.columns([2, 2, 2])

        with col1:
            if st.button("Generer un cv"):
                show_modifier_cv()

        with col3:
            if st.button("Generer une lettre de motivation personnalisée"):
                # Génération d'une lettre de motivation
                st.success("Génération de la lettre de motivation en cours...")
                lettre_motiv_personaliser = generate_lettre_motivation(
                    text_brut=st.session_state["cv_brut"],
                    job_offer=st.session_state["offre_emploi_brut"],
                )
                st.write(lettre_motiv_personaliser)

                doc = Document()
                doc.add_heading("Lettre de Motivation", 0)
                doc.add_paragraph(lettre_motiv_personaliser)

                # Sauvegarder en mémoire
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Bouton de téléchargement
                st.download_button(
                    label="📥 Télécharger la lettre en format Word",
                    data=buffer,
                    file_name="lettre_motivation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    with tab2:
        # Detailed comparison by section
        sections = {
            "Formation": ("formation", "🎓"),
            "Compétences": ("competences", "💪"),
            "Expériences": ("experiences", "💼"),
            "Profil": ("profil", "👤"),
        }

        for title, (key, emoji) in sections.items():
            with st.expander(
                f"{emoji} {title} - {(comparaison.get(key, 0) * 100):.1f}% de correspondance"
            ):
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("##### 📄 Votre CV")
                    st.markdown('<div class="section-content">', unsafe_allow_html=True)
                    if key.lower() == "competences":
                        st.write("• " + "\n• ".join(cv_data.get("Competences", [])))
                    elif key.lower() == "formation":
                        for formation in cv_data.get("Formation", []):
                            st.write(
                                f"• {formation.get('niveau_etudes')} - {', '.join(formation.get('domaine_etudes', []))}"
                            )
                    elif key.lower() == "experiences":
                        for exp in cv_data.get("Experiences", []):
                            st.write(
                                f"• {exp.get('poste_occupe')} ({exp.get('duree')})"
                            )
                    elif key.lower() == "profil":
                        st.write(f"• Titre: {cv_data.get('Profil', {}).get('titre')}")
                        st.write(
                            f"• Disponibilité: {cv_data.get('Profil', {}).get('disponibilite')}"
                        )
                    st.markdown("</div>", unsafe_allow_html=True)

                with col2:
                    st.markdown("##### 💼 Offre d'emploi")
                    st.markdown('<div class="section-content">', unsafe_allow_html=True)
                    if key.lower() == "competences":
                        st.write("• " + "\n• ".join(offre_data.get("Competences", [])))
                    elif key.lower() == "formation":
                        for formation in cv_data.get("Formation", []):
                            st.write(
                                f"• {offre_data.get('niveau_etudes')} - {', '.join(offre_data.get('domaine_etudes', []))}"
                            )
                    elif key.lower() == "experiences":
                        for exp in offre_data.get("Experiences", []):
                            st.write(
                                f"• {exp.get('poste_occupe')} ({exp.get('duree')})"
                            )
                    elif key.lower() == "profil":
                        st.write(
                            f"• Titre: {offre_data.get('Profil', {}).get('titre')}"
                        )
                        st.write(
                            f"• Disponibilité: {offre_data.get('Profil', {}).get('disponibilite')}"
                        )
                    st.markdown("</div>", unsafe_allow_html=True)

    # Overall recommendation
    total_score = sum(comparaison.values()) / len(comparaison)
    st.markdown("---")
    if total_score > 0.5:
        st.success(
            f"✨ Votre profil correspond bien à cette offre avec une compatibilité globale de {total_score*100:.1f}%"
        )
    else:
        st.warning(
            f"⚠️ Votre profil présente quelques écarts avec cette offre (compatibilité: {total_score*100:.1f}%)"
        )


@st.dialog("Domaine de l'offre", width="medium")
def entrer_domain():
    """
    Affiche une boîte de dialogue pour entrer le domaine de l'offre d'emploi.
    """
    st.markdown("### 📝 Domaine de l'offre d'emploi")
    st.markdown(
        "Veuillez entrer le domaine de l'offre d'emploi pour obtenir des recommandations personnalisées."
    )

    # Champ de saisie pour le domaine de l'offre d'emploi (obligatoire)
    domaine = st.text_input("Domaine de l'offre d'emploi", placeholder="Informatique")

    # Champ de saisie pour le département de l'offre d'emploi (optionnel mais doit être numérique)
    departement = st.text_input("Département de l'offre d'emploi", placeholder="69")

    # Bouton pour valider les informations
    if st.button("Valider"):
        # Validation du domaine (obligatoire)
        if not domaine:
            st.error("Veuillez entrer un domaine pour l'offre d'emploi.")
            return

        # Validation du département (optionnel mais doit être numérique)
        if departement and not departement.isdigit():
            st.error("Le département doit être un nombre valide.")
            return

        # Enregistrer les informations dans la session
        st.session_state["domaine_offre"] = domaine
        st.session_state["departement_offre"] = departement
        st.session_state["domaine_page"] = 2  # Passer à la page suivante

        # Fermer la boîte de dialogue et recharger la page
        st.rerun()


def afficher_carte_cliquable(offre_emploi, url="https://www.example.com"):
    """
    Affiche une carte cliquable stylisée pour une offre d'emploi.
    """
    # Custom CSS for job cards
    st.markdown(
        """
        <style>
            .job-card {
                background-color: white;
                border-radius: 10px;
                padding: 20px;
                margin: 10px 0;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                transition: transform 0.2s ease;
            }
            .job-card:hover {
                transform: translateY(-5px);
                box-shadow: 0 4px 8px rgba(0,0,0,0.2);
            }
            .job-title {
                color: #1E3D59;
                font-size: 20px;
                margin-bottom: 10px;
            }
            .job-info {
                color: #666;
                font-size: 16px;
                margin: 5px 0;
            }
            .competence-tag {
                background-color: #e9ecef;
                border-radius: 15px;
                padding: 5px 10px;
                margin: 2px;
                display: inline-block;
                font-size: 14px;
            }
            .match-button {
                background-color: #4CAF50;
                color: white;
                padding: 10px 20px;
                border-radius: 5px;
                border: none;
                cursor: pointer;
                margin-top: 15px;
            }
        </style>
    """,
        unsafe_allow_html=True,
    )

    with st.container():

        score = calculate_section_similarity(
            cv_reformule=json.dumps(st.session_state.get("cv_json", {})),
            offre_emploi_reformule=json.dumps(offre_emploi),
        )
        score_general = sum(score.values()) / len(score)
        st.markdown('<div class="job-card">', unsafe_allow_html=True)

        # En-tête de la carte
        profil = offre_emploi.get("Profil", {})
        st.markdown(
            f'<h3 class="job-title">👤 {profil.get("titre", "Titre non spécifié")}</h3>',
            unsafe_allow_html=True,
        )

        # Informations principales
        col1, col2 = st.columns([3, 1])

        with col1:
            st.markdown(
                f'<p class="job-info">📍 {profil.get("disponibilite", "Non spécifié")}</p>',
                unsafe_allow_html=True,
            )

            # Compétences
            st.markdown("#### 🛠️ Compétences requises")
            competences = offre_emploi.get("Competences", [])
            if competences:
                st.markdown(
                    f'<span class="competence-tag">{competences}</span>',
                    unsafe_allow_html=True,
                )
            else:
                st.info("Aucune compétence spécifiée.")

            data = score
            col11, col22 = st.columns([1, 1])
            with col11:
                # Affichage des scores avec st.metric
                st.metric(label="Formation", value=f"{data['formation']:.3f}")
                st.metric(label="Compétences", value=f"{data['competences']:.3f}")
            with col22:
                st.metric(label="Expériences", value=f"{data['experiences']:.3f}")
                st.metric(label="Profil", value=f"{data['profil']:.3f}")
            # st.write(score)
        with col2:
            # url = "https://www.example.com"

            # Afficher un bouton qui redirige vers le site externe
            st.markdown(
                f"""
                <a href="{url}" target="_blank">
                    <button style="background-color: #45a049; color: white; border-radius: 10px; padding: 10px 20px; border: none;">
                        Visiter le site externe
                    </button>
                </a>
            """,
                unsafe_allow_html=True,
            )

            st.metric(label="Score général", value=f"{score_general:.3f}")

            # if st.button("✨ Matching", key=f"match_{profil.get('titre', '')}"):
            #     st.session_state['offre_json'] = offre_emploi
            #     comparer_cv_offre()
            #     st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)


@st.dialog("offre_emploi", width="large")
def show_offre_emploi():
    """
    Affiche les offres d'emploi avec un design amélioré.
    """
    st.markdown(
        """
        <style>
            .search-header {
                background-color: #f8f9fa;
                padding: 20px;
                border-radius: 10px;
                margin-bottom: 20px;
            }
            .search-title {
                color: #1E3D59;
                font-size: 24px;
                margin-bottom: 10px;
            }
            .search-info {
                color: #666;
                font-size: 16px;
            }
        </style>
    """,
        unsafe_allow_html=True,
    )

    # En-tête de recherche
    st.markdown('<div class="search-header">', unsafe_allow_html=True)
    st.markdown(
        '<h2 class="search-title">🔍 Résultats de recherche</h2>',
        unsafe_allow_html=True,
    )

    domain_recherche = st.session_state.get("domaine_offre", "")
    departement_recherche = st.session_state.get("departement_offre", "")

    st.markdown(
        f'<p class="search-info">Domaine : {domain_recherche}</p>',
        unsafe_allow_html=True,
    )
    if departement_recherche:
        st.markdown(
            f'<p class="search-info">Département : {departement_recherche}</p>',
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

    # Affichage des offres avec barre de progression
    scraper = ScrapingFactory()
    with st.spinner("Recherche des offres en cours..."):
        domain_recherche = st.session_state.get('domaine_offre', "")

        offres = scraper.scrap_many('domaine_recherche', limit=2)
        
        progress_bar = st.progress(0)
        liste_offres_analyser = []

        for i, offre in enumerate(offres):
            progress = (i + 1) / len(offres)
            progress_bar.progress(progress)

            offre_text = f"{offre['description']} "
            url = offre["origineOffre"]["urlOrigine"]
            # st.write(url)
            analyse = analyze_offre_emploi(offre_text)
            analyse = json.loads(analyse)

            afficher_carte_cliquable(analyse, url)
            liste_offres_analyser.append(analyse)
            # time.sleep(2)

        progress_bar.empty()


@st.dialog("Modification du CV", width="large")
def show_modifier_cv():
    """
    Affiche un dialogue pour modifier les informations du CV.
    cv_info peut être une chaîne JSON ou un dictionnaire Python.
    """

    cv_info = st.session_state.get(
        "cv_json",
        {
            "Profil": {"titre": "", "disponibilite": ""},
            "Formation": [{"niveau_etudes": "", "domaine_etudes": []}],
            "Competences": [],
            "Experiences": [],
        },
    )
    # Si cv_info est une chaîne JSON, la convertir en dictionnaire
    if isinstance(cv_info, str):
        try:
            cv_info = json.loads(cv_info)
        except json.JSONDecodeError:
            st.error("Le format du CV est invalide. Veuillez charger un CV valide.")
            return

    # Vérifier que cv_info est un dictionnaire
    if not isinstance(cv_info, dict):
        st.error(
            "Les données du CV doivent être un dictionnaire ou une chaîne JSON valide."
        )
        return

    st.markdown("### 📝 Visualisation et Modification du CSV")
    st.markdown("---")

    # Section Profil
    with st.expander("👤 Profil", expanded=True):
        st.markdown("#### Profil")
        cv_info["Profil"]["titre"] = st.text_input(
            "Titre", cv_info["Profil"].get("titre", "")
        )
        cv_info["Profil"]["disponibilite"] = st.text_input(
            "Disponibilité", cv_info["Profil"].get("disponibilite", "")
        )

    # Section Compétences
    with st.expander("🛠️ Compétences", expanded=False):
        st.markdown("#### Compétences")
        competences_str = st.text_area(
            "Liste des compétences (séparées par des virgules)",
            ", ".join(cv_info.get("Competences", [])),
        )
        cv_info["Competences"] = [c.strip() for c in competences_str.split(",")]

    # Section Expériences
    with st.expander("💼 Expériences", expanded=False):
        st.markdown("#### Expériences")
        for i, experience in enumerate(cv_info.get("Experiences", [])):
            st.markdown(f"##### Expérience {i + 1}")
            experience["poste_occupe"] = st.text_input(
                "Poste occupé",
                experience.get("poste_occupe", ""),
                key=f"poste_occupe_{i}",
            )
            experience["domaine_activite"] = st.text_input(
                "Domaine d'activité (séparés par des virgules)",
                ", ".join(experience.get("domaine_activite", [])),
                key=f"domaine_activite_{i}",
            )
            experience["duree"] = st.text_input(
                "Durée", experience.get("duree", ""), key=f"duree_{i}"
            )

    # Section Formations
    with st.expander("🎓 Formations", expanded=False):
        st.markdown("#### Formations")
        for i, formation in enumerate(cv_info.get("Formation", [])):
            st.markdown(f"##### Formation {i + 1}")
            formation["niveau_etudes"] = st.text_input(
                "Niveau d'études",
                formation.get("niveau_etudes", ""),
                key=f"niveau_etudes_{i}",
            )
            formation["domaine_etudes"] = st.text_input(
                "Domaine d'études (séparés par des virgules)",
                ", ".join(formation.get("domaine_etudes", [])),
                key=f"domaine_etudes_{i}",
            )

    # Boutons de validation et d'annulation
    col1, temp, col2 = st.columns([2, 3, 1])
    with col1:
        if st.button("Enregistrer les modifications", type="primary"):
            st.session_state["cv_json"] = (
                cv_info  # Mettre à jour les informations du CV dans la session
            )
            st.success("Les modifications ont été enregistrées avec succès !")
            st.session_state["modifications"] = True
            # st.session_state["suivant"] = True

            st.rerun()  # Recharger la page pour appliquer les modifications
